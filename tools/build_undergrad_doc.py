from __future__ import annotations

import shutil
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "chapter 1-4 updated diox.docx"
BACKUP_PATH = ROOT / "chapter 1-4 updated diox.backup.docx"
PLACEHOLDER_DIR = ROOT / "outputs" / "doc_placeholders"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    hdr = table.rows[0].cells
    for cell, header in zip(hdr, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    doc.add_paragraph()


def add_para(doc: Document, text: str, style: str | None = None, align=None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def create_placeholder(name: str, title: str) -> Path:
    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    path = PLACEHOLDER_DIR / f"{name}.png"
    fig = plt.figure(figsize=(7.2, 3.8), dpi=160)
    ax = fig.add_subplot(111)
    ax.set_facecolor("#F4F6F9")
    for spine in ax.spines.values():
        spine.set_edgecolor("#AAB4C0")
        spine.set_linewidth(1.8)
    ax.set_xticks([])
    ax.set_yticks([])
    ax.text(0.5, 0.58, "PLACEHOLDER IMAGE", ha="center", va="center", fontsize=18, color="#44546A", weight="bold")
    ax.text(0.5, 0.42, title, ha="center", va="center", fontsize=11, color="#44546A", wrap=True)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)
    return path


def add_figure_placeholder(doc: Document, name: str, title: str, caption: str) -> None:
    path = create_placeholder(name, title)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(5.9))
    doc.inline_shapes[-1]._inline.docPr.set("descr", caption)
    doc.inline_shapes[-1]._inline.docPr.set("title", title)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.paragraph_format.space_after = Pt(10)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def fmt(value: float, digits: int = 2) -> str:
    return f"{value:,.{digits}f}"


def collect_results() -> dict[str, object]:
    data = pd.read_csv(ROOT / "data/processed/modeling_dataset.csv")
    metrics = pd.read_csv(ROOT / "outputs/metrics/model_comparison.csv")
    validation = pd.read_csv(ROOT / "outputs/metrics/validation_predictions.csv")
    forecast = pd.read_csv(ROOT / "outputs/forecasts/latest_forecast.csv")
    best = metrics[metrics["model"] == "xgboost"].iloc[0]
    latest = data.sort_values(["year", "epi_week"]).groupby("state").tail(1)
    return {
        "data": data,
        "metrics": metrics,
        "validation": validation,
        "forecast": forecast,
        "best": best,
        "latest": latest,
        "rows": len(data),
        "states": data["state"].nunique(),
        "periods": data[["year", "epi_week"]].drop_duplicates().shape[0],
        "year_min": int(data["year"].min()),
        "year_max": int(data["year"].max()),
        "cases": int(data["suspected_cases"].sum()),
        "deaths": int(data["deaths"].fillna(0).sum()),
        "cfr": float(data["deaths"].sum() / data["suspected_cases"].sum()),
        "risk_counts": data["risk_level"].value_counts().to_dict(),
        "latest_risk_counts": latest["risk_level"].value_counts().to_dict(),
    }


def add_title_page(doc: Document) -> None:
    for text, size, bold in [
        ("DEVELOPMENT OF A PREDICTIVE MODEL FOR CHOLERA OUTBREAKS IN NIGERIA", 16, True),
        ("BY", 12, True),
        ("ADEBIYI AL-MUSTAPHA EMMANUEL", 14, True),
        ("MATRIC NUMBER: 2022/42606", 12, True),
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        paragraph.paragraph_format.space_after = Pt(16)
    for text in [
        "A PROJECT SUBMITTED TO THE DEPARTMENT OF SOFTWARE ENGINEERING, FACULTY OF COMPUTING AND INFORMATION TECHNOLOGY, OSUN STATE UNIVERSITY, OSUN STATE, NIGERIA.",
        "IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF BACHELOR OF SCIENCE DEGREE IN SOFTWARE ENGINEERING",
        "SUPERVISOR: DR T. A. ADEKUNLE",
        "MARCH 2026",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.size = Pt(12)
        if text.startswith("SUPERVISOR") or text == "MARCH 2026":
            run.bold = True
        paragraph.paragraph_format.space_after = Pt(14)
    doc.add_page_break()


def add_abstract(doc: Document, r: dict[str, object]) -> None:
    doc.add_heading("ABSTRACT", level=1)
    add_para(
        doc,
        "Cholera remains a major public health concern in Nigeria because outbreaks are influenced by delayed reporting, sanitation conditions, rainfall, flooding, and movement across communities. This project developed a weekly state-level cholera forecasting and risk dashboard for Nigeria using reported cholera cases and deaths, epidemiological week information, lagged case features, environmental variables, and machine learning models. The system was implemented as a modular Python and React application consisting of PDF extraction, raw data processing, feature engineering, model training, FastAPI model serving, and a React dashboard.",
    )
    add_para(
        doc,
        f"The processed modelling dataset contained {r['rows']} state-period records across {r['states']} Nigerian states and the Federal Capital Territory, covering {r['year_min']} to {r['year_max']}. The dataset recorded {r['cases']:,} suspected cases and {r['deaths']:,} deaths, giving an overall case fatality ratio of {r['cfr'] * 100:.2f}%. Risk classes were derived from case-count thresholds after trimming high outliers using an interquartile range method. The final dataset contained {r['risk_counts'].get('High', 0)} high-risk records, {r['risk_counts'].get('Medium', 0)} medium-risk records, and {r['risk_counts'].get('Low', 0)} low-risk records.",
    )
    best = r["best"]
    add_para(
        doc,
        f"Five forecasting approaches were compared using expanding time-based cross-validation: naive lag-1 baseline, four-period moving average, Random Forest, XGBoost, and Prophet. XGBoost was selected as the best trainable model because it produced the lowest cross-validation RMSE among the machine learning models, with MAE of {fmt(best['mae'])}, RMSE of {fmt(best['rmse'])}, SMAPE of {fmt(best['smape'])}%, and R2 of {fmt(best['r2'], 3)}. A four-week recursive forecast was generated for each state and converted into low, medium, and high risk labels. The system also included a React dashboard with a Nigeria state choropleth map, KPI cards, epi-week filtering, case trend chart, forecast chart with uncertainty interval, and a state risk summary page. The project demonstrates that a data-driven decision-support prototype can help summarize cholera risk patterns, although it should not replace verified public health surveillance or expert epidemiological judgement.",
    )
    doc.add_page_break()


def add_chapter_one(doc: Document) -> None:
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("INTRODUCTION", level=1)
    doc.add_heading("1.1 Background of the Study", level=2)
    add_para(doc, "Cholera is an acute diarrhoeal disease caused by ingestion of food or water contaminated with Vibrio cholerae. It spreads rapidly where access to safe water, sanitation, and hygiene is poor. In Nigeria, cholera outbreaks often occur in relation to seasonal rainfall, flooding, population movement, poor drainage, and delays in health surveillance reporting. Because the disease can lead to severe dehydration and death within a short time, early detection and response remain important public health priorities.")
    add_para(doc, "Traditional cholera response depends mainly on confirmed reports after suspected cases have already been recorded. This approach is useful for treatment and control, but it is reactive. A predictive system can support preparedness by analysing historical case patterns, recent reporting periods, environmental conditions, and spatial distribution across states. Such a system can help public health officers identify areas that may require closer monitoring, water sanitation intervention, or medical resource allocation.")
    add_para(doc, "This project focuses on developing a weekly state-level predictive model for cholera outbreak risk in Nigeria. The implemented system uses historical NCDC-style cholera reporting data, case and death counts, case fatality ratio, epidemiological week ranges, climate variables, lagged cases, and rolling statistics. The prediction output is displayed through a React dashboard and served through a FastAPI backend.")
    doc.add_heading("1.2 Statement of the Problem", level=2)
    add_para(doc, "Cholera surveillance reports are often irregular. Some reports are weekly, while others summarize several epidemiological weeks, such as weeks 1-4 or 6-9. This irregular reporting pattern makes direct time-series forecasting difficult. In addition, outbreak data are highly skewed because a few state-weeks may contain very large spikes while many others contain low case counts. If such outliers are not handled properly, risk thresholds and model evaluation may become misleading.")
    add_para(doc, "Another problem is that public health data are usually stored as reports, spreadsheets, or PDFs rather than as a ready-to-use modelling dataset. This creates a need for a pipeline that can separate raw data collection from feature engineering, so that updated NCDC data can replace temporary data without rewriting the model code. There is also a need to compare more than one algorithm instead of relying on a single model.")
    doc.add_heading("1.3 Aim and Objectives", level=2)
    add_para(doc, "The aim of this study is to develop a predictive model and dashboard for forecasting state-level cholera outbreak risk in Nigeria.")
    add_bullets(doc, [
        "To prepare a raw-data pipeline that can accept NCDC cholera case and death data as a swappable CSV file.",
        "To extract and engineer features such as epidemiological week bounds, reporting gaps, lagged cases, rolling averages, rainfall, temperature, humidity, and case fatality ratio.",
        "To train and compare baseline, Random Forest, XGBoost, and Prophet models for short-term cholera case forecasting.",
        "To classify predicted cases into low, medium, and high risk using thresholds derived after outlier trimming.",
        "To serve model predictions through FastAPI routes and visualize results in a React dashboard.",
    ])
    doc.add_heading("1.4 Research Questions", level=2)
    add_numbered(doc, [
        "How can irregular weekly and multi-week cholera situation report data be transformed into a modelling dataset?",
        "Which forecasting approach performs best among baseline methods, Random Forest, XGBoost, and Prophet on the available data?",
        "How can forecasted case counts be converted into interpretable public health risk levels?",
        "How can model outputs be presented in an understandable dashboard for decision support?",
    ])
    doc.add_heading("1.5 Scope of the Study", level=2)
    add_para(doc, "The study is limited to weekly state-level cholera outbreak forecasting for Nigeria. It uses cases, deaths, case fatality ratio, epidemiological week values, state identifiers, lag features, rolling features, and environmental variables where available. The dashboard provides a Nigeria map, KPIs, trends, forecasts, and a state risk summary. It does not replace official NCDC reporting, laboratory confirmation, clinical diagnosis, or decisions by qualified public health professionals.")
    doc.add_heading("1.6 Significance of the Study", level=2)
    add_para(doc, "The project is significant because it demonstrates how software engineering and machine learning can support public health surveillance. It provides a reproducible workflow from raw data to dashboard visualization and shows how model outputs can be communicated using risk categories and forecast charts. For an undergraduate software engineering project, it also demonstrates modular design, API serving, data processing, and frontend implementation.")
    doc.add_heading("1.7 Definition of Key Terms", level=2)
    add_bullets(doc, [
        "Epidemiological week: A standardized week numbering system used in disease surveillance reporting.",
        "Case fatality ratio: The proportion of deaths among reported cases, calculated as deaths divided by cases.",
        "Feature engineering: The process of transforming raw data into variables suitable for machine learning.",
        "Forecasting: The prediction of future values based on historical and recent data.",
        "Risk classification: The conversion of predicted case counts into categories such as low, medium, and high.",
    ])
    doc.add_page_break()


def add_chapter_two(doc: Document) -> None:
    doc.add_heading("CHAPTER TWO", level=1)
    doc.add_heading("LITERATURE REVIEW", level=1)
    doc.add_heading("2.1 Introduction", level=2)
    add_para(doc, "This chapter reviews the major concepts that support the development of a predictive model for cholera outbreaks. The review covers cholera outbreak surveillance, predictive modelling in public health, environmental factors, machine learning models, time-series forecasting, and dashboard-based decision support.")
    doc.add_heading("2.2 Cholera Outbreak Surveillance", level=2)
    add_para(doc, "Cholera surveillance involves collecting, organizing, and interpreting case reports from health facilities and public health agencies. Surveillance data usually contain information such as location, number of suspected cases, deaths, and reporting period. In Nigeria, cholera situation reports and weekly epidemiological reports are important sources for tracking outbreak patterns across states.")
    add_para(doc, "A common challenge in outbreak surveillance is irregular reporting. Some situation reports cover one week, while others summarize multiple weeks. This affects modelling because the reporting period may not represent a uniform time interval. Therefore, a robust system must preserve the original reporting label while extracting start week, end week, and period length for feature engineering.")
    doc.add_heading("2.3 Environmental Factors and Cholera", level=2)
    add_para(doc, "Cholera transmission is linked with environmental and water-related factors. Heavy rainfall can contaminate water sources, flooding can spread waste into drinking water supplies, and high humidity may support environmental persistence. For this reason, environmental variables such as rainfall, temperature, and humidity can improve outbreak forecasting when they are available and properly aligned with epidemiological weeks.")
    doc.add_heading("2.4 Machine Learning in Disease Prediction", level=2)
    add_para(doc, "Machine learning is useful in disease prediction because it can identify nonlinear relationships among variables. Random Forest is an ensemble method that combines many decision trees and is suitable for structured tabular data. XGBoost is a gradient boosting method that often performs well on structured datasets because it sequentially improves weak learners. Prophet is a time-series forecasting tool that models trend and seasonality, making it useful for outbreak data with temporal patterns.")
    doc.add_heading("2.5 Related Empirical Studies", level=2)
    add_para(doc, "Previous studies have applied machine learning and time-series methods to infectious disease prediction. Some studies used Random Forest and Support Vector Machines for cholera prediction, while others combined climate variables and surveillance data. Geographic information systems have also been used to map disease risk and support outbreak response. These studies show that historical disease data, environmental variables, and spatial visualization can improve public health preparedness.")
    doc.add_heading("2.6 Gap in Existing Work", level=2)
    add_para(doc, "Many existing systems either focus only on modelling or only on visualization. Some also assume regular reporting intervals, which may not reflect real situation report data. This project addresses the gap by building a modular end-to-end prototype that supports irregular reporting periods, multiple model comparison, robust risk thresholds, API model serving, and a React dashboard.")
    doc.add_page_break()


def add_chapter_three(doc: Document, r: dict[str, object]) -> None:
    doc.add_heading("CHAPTER THREE", level=1)
    doc.add_heading("METHODOLOGY AND SYSTEM DESIGN", level=1)
    doc.add_heading("3.1 Introduction", level=2)
    add_para(doc, "This chapter describes the research design, data processing workflow, feature engineering, model development, evaluation strategy, API design, and dashboard implementation used in the project.")
    doc.add_heading("3.2 Research Design", level=2)
    add_para(doc, "The study adopted an experimental and data-driven software engineering design. Historical cholera case data were transformed into a modelling dataset, several forecasting algorithms were trained and evaluated, and the selected model was served through an API for dashboard consumption.")
    doc.add_heading("3.3 System Architecture", level=2)
    add_para(doc, "The system was implemented as separate modules. The PDF extraction module downloads and extracts NCDC cholera situation report tables. The data pipeline loads the raw cholera dataset, normalizes state names and columns, computes CFR where necessary, and merges climate variables. The feature engineering module creates time, lag, rolling, reporting-gap, and risk-label features. The training module compares models and saves the best model. The FastAPI service exposes model and dashboard routes. The React dashboard presents the map, KPIs, charts, forecast, and state risk summary.")
    add_figure_placeholder(doc, "figure_3_1_architecture", "System architecture screenshot/diagram to be inserted", "Figure 3.1: Architecture of the cholera forecasting and dashboard system.")
    doc.add_heading("3.4 Data Collection and Preparation", level=2)
    add_para(doc, f"The active processed dataset contained {r['rows']} records from {r['states']} states including the Federal Capital Territory. It covered {r['year_min']} to {r['year_max']} and included {r['periods']} unique reporting periods. The raw cholera dataset is stored as cholera_data.csv so that manually collected NCDC data can later replace the temporary dataset without changing the pipeline.")
    add_para(doc, "The raw dataset contains state, year, epidemiological week, cases, deaths, and optional CFR. The system accepts both single-week values and multi-week ranges such as 1-4 and 6-9. State names are normalized to support joins with the Nigeria boundary GeoJSON used by the dashboard.")
    doc.add_heading("3.5 Feature Engineering", level=2)
    add_para(doc, "The feature engineering process created epi_week_start, epi_week_end, period_weeks, state_code, report_gap_weeks, date, month, quarter, rainy_season, rainfall_mm, temperature_c, humidity_pct, lag_1_cases, lag_2_cases, lag_4_cases, lag_8_cases, lagged deaths, lagged CFR, rolling case averages, rolling death averages, and rolling CFR averages. These features allow the model to learn both temporal and state-level patterns.")
    add_para(doc, "Risk labels were derived from suspected cases using median and seventy-fifth percentile thresholds after trimming extreme high outliers with the interquartile range upper fence. This made the low, medium, and high risk classes more representative of the main data distribution.")
    doc.add_heading("3.6 Model Development", level=2)
    add_para(doc, "Five approaches were compared: naive lag-1 baseline, four-period moving average, Random Forest Regressor, XGBoost Regressor, and Prophet. The machine learning models used lagged case features, reporting-period features, climate variables, and CFR-related features. Forecasts were generated recursively for four weeks per state.")
    add_table(doc, ["Model", "Purpose"], [
        ["Naive lag-1", "Baseline using previous case count as forecast"],
        ["Moving average", "Baseline using recent rolling case average"],
        ["Random Forest", "Tree ensemble model for structured tabular data"],
        ["XGBoost", "Gradient boosting model for nonlinear structured data"],
        ["Prophet", "Time-series model for trend and seasonality comparison"],
    ])
    doc.add_heading("3.7 Data Splitting and Evaluation", level=2)
    add_para(doc, "Because the dataset is time ordered, random train-test splitting was not used. The data were sorted chronologically from 2021 to 2025. Model comparison used three expanding time-based cross-validation folds. The latest holdout split trained on earlier periods and tested on future periods, ensuring that 2025 records were treated as future data rather than past data.")
    add_para(doc, "The evaluation metrics used were Mean Absolute Error, Root Mean Squared Error, Symmetric Mean Absolute Percentage Error, and the coefficient of determination.")
    doc.add_heading("3.8 API and Dashboard Implementation", level=2)
    add_para(doc, "The backend was implemented with FastAPI. The main routes include health check, model status, summary, history, forecast, boundaries, predict, batch predict, metrics, and artifact reload. The frontend was implemented with React, Vite, and Leaflet. The dashboard contains two pages: Overview and State Risk Summary.")
    add_figure_placeholder(doc, "figure_3_2_dashboard_placeholder", "React dashboard screenshot to be inserted", "Figure 3.2: Placeholder for the React dashboard overview page.")
    doc.add_page_break()


def add_chapter_four(doc: Document, r: dict[str, object]) -> None:
    data = r["data"]
    metrics = r["metrics"]
    validation = r["validation"]
    forecast = r["forecast"]
    latest = r["latest"]
    doc.add_heading("CHAPTER FOUR", level=1)
    doc.add_heading("RESULTS AND DISCUSSION", level=1)
    doc.add_heading("4.1 Introduction", level=2)
    add_para(doc, "This chapter presents the implementation results, dataset summary, model comparison results, validation results, forecast output, risk classification output, dashboard results, and discussion of findings.")
    doc.add_heading("4.2 Dataset Summary", level=2)
    add_para(doc, f"The processed dataset contained {r['rows']} state-period observations, {r['states']} states, and {r['periods']} distinct reporting periods from {r['year_min']} to {r['year_max']}. A total of {r['cases']:,} suspected cases and {r['deaths']:,} deaths were recorded, giving an overall CFR of {r['cfr'] * 100:.2f}%.")
    add_table(doc, ["Statistic", "Value"], [
        ["Processed records", f"{r['rows']:,}"],
        ["States covered", r["states"]],
        ["Reporting periods", r["periods"]],
        ["Year range", f"{r['year_min']} - {r['year_max']}"],
        ["Total suspected cases", f"{r['cases']:,}"],
        ["Total deaths", f"{r['deaths']:,}"],
        ["Overall CFR", f"{r['cfr'] * 100:.2f}%"],
    ])
    add_para(doc, f"The risk distribution in the modelling dataset was {r['risk_counts'].get('Low', 0)} low-risk records, {r['risk_counts'].get('Medium', 0)} medium-risk records, and {r['risk_counts'].get('High', 0)} high-risk records. In the latest available record per state, {r['latest_risk_counts'].get('Low', 0)} states were low risk, {r['latest_risk_counts'].get('Medium', 0)} were medium risk, and {r['latest_risk_counts'].get('High', 0)} were high risk.")
    top_latest = latest.sort_values("suspected_cases", ascending=False)[["state", "year", "epi_week_label", "suspected_cases", "deaths", "cfr", "risk_level"]].head(10)
    add_table(doc, ["State", "Year", "Epi week", "Cases", "Deaths", "CFR", "Risk"], [
        [row.state, int(row.year), row.epi_week_label, int(row.suspected_cases), int(row.deaths), f"{row.cfr * 100:.2f}%", row.risk_level]
        for row in top_latest.itertuples()
    ])
    add_figure_placeholder(doc, "figure_4_1_map", "Risk map screenshot to be inserted", "Figure 4.1: Placeholder for the Nigeria choropleth risk map.")
    doc.add_heading("4.3 Model Comparison Results", level=2)
    add_para(doc, "The models were evaluated using three expanding time-based cross-validation folds. This approach ensured that earlier reporting periods were used for training and later periods were used for testing. The comparison included two baselines and three modelling approaches.")
    add_table(doc, ["Model", "Status", "Folds", "MAE", "RMSE", "SMAPE", "R2"], [
        [row.model, row.status, int(row.folds), fmt(row.mae), fmt(row.rmse), f"{fmt(row.smape)}%", fmt(row.r2, 3)]
        for row in metrics.itertuples()
    ])
    best = r["best"]
    add_para(doc, f"XGBoost was selected as the best model because it produced the lowest RMSE among the trainable machine learning models, with MAE of {fmt(best['mae'])}, RMSE of {fmt(best['rmse'])}, SMAPE of {fmt(best['smape'])}%, and R2 of {fmt(best['r2'], 3)}. Random Forest produced a slightly lower MAE of {fmt(metrics[metrics['model'] == 'random_forest'].iloc[0]['mae'])}, but its RMSE was higher than XGBoost. Prophet performed poorly on the current irregular state-level dataset, which suggests that the available data are not yet strong enough for Prophet to generalize reliably across state reporting patterns.")
    doc.add_heading("4.4 Validation Result", level=2)
    add_para(doc, f"The latest holdout validation set contained {len(validation)} records. The mean actual case count was {validation['actual_cases'].mean():.2f}, while the mean predicted case count was {validation['predicted_cases'].mean():.2f}. The mean absolute error on the holdout validation predictions was {validation['absolute_error'].mean():.2f}. The largest error occurred in Bayelsa in 2025 week 4, where actual cases were 605 while the model predicted 14.15 cases. This shows that sudden outbreak spikes remain difficult to predict with the current feature set.")
    pred_risk = validation["predicted_risk_level"].value_counts().to_dict()
    add_table(doc, ["Predicted risk level", "Validation count"], [
        ["Low", pred_risk.get("Low", 0)],
        ["Medium", pred_risk.get("Medium", 0)],
        ["High", pred_risk.get("High", 0)],
    ])
    add_figure_placeholder(doc, "figure_4_2_trend", "Case trend and moving average screenshot to be inserted", "Figure 4.2: Placeholder for case trend visualization.")
    doc.add_heading("4.5 Forecast Results", level=2)
    forecast_counts = forecast["risk_level"].value_counts().to_dict()
    add_para(doc, f"The trained model generated four-week recursive forecasts for each state, producing {len(forecast)} forecast records. The forecast output contained {forecast_counts.get('Low', 0)} low-risk forecasts, {forecast_counts.get('Medium', 0)} medium-risk forecasts, and {forecast_counts.get('High', 0)} high-risk forecasts. Confidence-style intervals were added using recent state-level historical volatility.")
    top_forecast = forecast.sort_values("predicted_cases", ascending=False).head(10)
    add_table(doc, ["State", "Forecast week", "Year", "Epi week", "Predicted cases", "Lower", "Upper", "Risk"], [
        [row.state, int(row.forecast_week), int(row.year), int(row.epi_week), fmt(row.predicted_cases), fmt(row.predicted_lower), fmt(row.predicted_upper), row.risk_level]
        for row in top_forecast.itertuples()
    ])
    add_figure_placeholder(doc, "figure_4_3_forecast", "Forecast chart with confidence interval screenshot to be inserted", "Figure 4.3: Placeholder for forecast chart with uncertainty interval.")
    doc.add_heading("4.6 Dashboard Results", level=2)
    add_para(doc, "The React dashboard was implemented with two pages. The Overview page displays KPI cards, a Nigeria state choropleth map, selected-state details, a case trend chart, and a four-week forecast chart. The State Risk Summary page displays the top ten states for the selected reporting period with risk level, cases, deaths, CFR, and a sparkline trend. The dashboard uses a top-right epi-week filter, and clicking a state on the map updates the state-specific forecast panel.")
    add_figure_placeholder(doc, "figure_4_4_summary", "State risk summary page screenshot to be inserted", "Figure 4.4: Placeholder for state risk summary dashboard page.")
    doc.add_heading("4.7 Discussion of Findings", level=2)
    add_para(doc, "The results show that lagged case features and tree-based models are useful for short-term state-level cholera forecasting. However, the negative R2 values indicate that the model still struggles to explain large variations in future case counts. This is expected because cholera data contain sudden spikes, reporting gaps, and multi-week aggregation. The outlier analysis confirms that major outbreak spikes, such as the Bayelsa 2025 week 4 record, remain the hardest cases to predict.")
    add_para(doc, "The dashboard results are useful for interpretation because they translate model outputs into risk classes, maps, and simple trend charts. The system is therefore best understood as an academic decision-support prototype rather than a fully operational public health warning system.")
    doc.add_heading("4.8 Limitations", level=2)
    add_bullets(doc, [
        "The current dataset is temporary and will be replaced with manually validated NCDC data.",
        "Some reports cover multiple epidemiological weeks, which reduces temporal precision.",
        "Future rainfall, temperature, and humidity values are approximated using recent averages.",
        "Very large outbreak spikes remain difficult to forecast.",
        "The dashboard is a prototype and should not be used as the sole basis for public health action.",
    ])
    doc.add_heading("4.9 Summary", level=2)
    add_para(doc, "This chapter presented the project results using actual system outputs. The dataset contained 604 processed observations across 37 states from 2021 to 2025. XGBoost was selected after expanding time-based cross-validation, and the dashboard successfully displayed map-based risk, KPI summaries, trends, forecasts, and state risk summaries.")


def add_references(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("REFERENCES", level=1)
    refs = [
        "Nigeria Centre for Disease Control and Prevention. Cholera situation reports and weekly epidemiological reports.",
        "NASA POWER Data Services. Climate data API documentation.",
        "Taylor, S. J., and Letham, B. Forecasting at scale. The American Statistician.",
        "Chen, T., and Guestrin, C. XGBoost: A scalable tree boosting system.",
        "Scikit-learn Developers. Random Forest Regressor and model evaluation documentation.",
        "World Health Organization. Cholera fact sheets and outbreak response guidance.",
    ]
    for ref in refs:
        add_para(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT)


def main() -> None:
    if DOC_PATH.exists() and not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)
    r = collect_results()
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_abstract(doc, r)
    add_chapter_one(doc)
    add_chapter_two(doc)
    add_chapter_three(doc, r)
    add_chapter_four(doc, r)
    add_references(doc)
    doc.save(DOC_PATH)
    print(f"Wrote {DOC_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
