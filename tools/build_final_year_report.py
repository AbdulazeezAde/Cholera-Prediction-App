from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "chapter 1-4 updated diox.docx"
PLACEHOLDER_DIR = ROOT / "outputs" / "doc_placeholders"


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char)

    text = OxmlElement("w:t")
    text.text = "Update field in Microsoft Word"
    run._r.append(text)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    pg_num_type.set(qn("w:start"), str(start))


def add_page_number(section, roman: bool = False) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    for paragraph in footer.paragraphs:
        paragraph.clear()
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE \\* ROMAN" if roman else "PAGE")


def clear_doc(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_font(run, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.bold = True

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.line_spacing = 1
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)


def paragraph(doc: Document, text: str = "", style: str | None = None, align=None, bold=False) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold)


def heading(doc: Document, text: str, level: int = 1, center: bool = False) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, bold=True)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run)


def number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, 12, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, caption: str, headers: list[str], rows: list[list[object]]) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_font(run, 10, italic=True)

    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)
    for cell, value in zip(t.rows[0].cells, headers):
        set_cell(cell, value, True)
        shade_cell(cell, "EDEDED")
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            set_cell(cell, value)
    paragraph(doc)


def placeholder_image(name: str, title: str) -> Path:
    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    path = PLACEHOLDER_DIR / f"{name}.png"
    fig = plt.figure(figsize=(6.8, 3.5), dpi=160)
    ax = fig.add_subplot(111)
    ax.set_xticks([])
    ax.set_yticks([])
    ax.set_facecolor("#f2f2f2")
    for spine in ax.spines.values():
        spine.set_color("#777777")
        spine.set_linewidth(1.5)
    ax.text(0.5, 0.58, "PLACEHOLDER IMAGE", ha="center", va="center", fontsize=17, weight="bold")
    ax.text(0.5, 0.42, title, ha="center", va="center", fontsize=10, wrap=True)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)
    return path


def figure(doc: Document, name: str, title: str, caption: str) -> None:
    path = placeholder_image(name, title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.8))
    doc.inline_shapes[-1]._inline.docPr.set("descr", caption)
    doc.inline_shapes[-1]._inline.docPr.set("title", title)
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(caption)
    set_font(cr, 10, italic=True)
    paragraph(doc)


def collect() -> dict[str, object]:
    data = pd.read_csv(ROOT / "data/processed/modeling_dataset.csv")
    metrics = pd.read_csv(ROOT / "outputs/metrics/model_comparison.csv")
    validation = pd.read_csv(ROOT / "outputs/metrics/validation_predictions.csv")
    forecast = pd.read_csv(ROOT / "outputs/forecasts/latest_forecast.csv")
    latest = data.sort_values(["year", "epi_week"]).groupby("state").tail(1)
    enrichment_files = {
        "NASA POWER climate": ROOT / "data/interim/climate_state_week.csv",
        "Rainfall flood proxies": ROOT / "data/interim/flood_state_week.csv",
        "World Bank WDI WASH": ROOT / "data/interim/wash_wdi_nigeria.csv",
        "State WASHNORM WASH": ROOT / "data/interim/wash_state_year.csv",
        "UNHCR/DTM displacement": ROOT / "data/interim/displacement_state_year.csv",
        "GRID3 health facilities": ROOT / "data/interim/health_facility_state.csv",
    }
    enrichment_status = []
    for name, path in enrichment_files.items():
        rows = len(pd.read_csv(path)) if path.exists() else 0
        enrichment_status.append({"source": name, "file": path.relative_to(ROOT).as_posix(), "rows": rows})
    return {
        "data": data,
        "metrics": metrics,
        "validation": validation,
        "forecast": forecast,
        "latest": latest,
        "enrichment_status": enrichment_status,
        "rows": len(data),
        "columns": len(data.columns),
        "states": data["state"].nunique(),
        "periods": data[["year", "epi_week"]].drop_duplicates().shape[0],
        "year_min": int(data["year"].min()),
        "year_max": int(data["year"].max()),
        "cases": int(data["suspected_cases"].sum()),
        "deaths": int(data["deaths"].fillna(0).sum()),
        "cfr": float(data["deaths"].sum() / data["suspected_cases"].sum()),
    }


def add_front_matter(doc: Document, r: dict[str, object]) -> None:
    set_page_number_format(doc.sections[0], "lowerRoman", 1)
    add_page_number(doc.sections[0], roman=True)

    for text in [
        "DEVELOPMENT OF A PREDICTIVE MODEL FOR CHOLERA OUTBREAKS IN NIGERIA",
        "BY",
        "ADEBIYI AL-MUSTAPHA EMMANUEL",
        "MATRIC NUMBER: 2022/42606",
    ]:
        paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc)
    paragraph(doc, "A PROJECT SUBMITTED TO THE DEPARTMENT OF SOFTWARE ENGINEERING, FACULTY OF COMPUTING AND INFORMATION TECHNOLOGY, OSUN STATE UNIVERSITY, OSUN STATE, NIGERIA.", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc, "IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF BACHELOR OF SCIENCE DEGREE IN SOFTWARE ENGINEERING", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc)
    paragraph(doc, "SUPERVISOR: DR T. A. ADEKUNLE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    paragraph(doc)
    paragraph(doc, "MARCH 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_page_break()

    heading(doc, "CERTIFICATION", center=True)
    paragraph(doc, "This is to certify that this project titled DEVELOPMENT OF A PREDICTIVE MODEL FOR CHOLERA OUTBREAKS IN NIGERIA was carried out by ADEBIYI AL-MUSTAPHA EMMANUEL with matriculation number 2022/42606 in the Department of Software Engineering, Faculty of Computing and Information Technology, Osun State University, Osun State, Nigeria.")
    paragraph(doc)
    paragraph(doc, "DR T. A. ADEKUNLE\t\t\t\t____________________________")
    paragraph(doc, "Project Supervisor")
    paragraph(doc)
    paragraph(doc, "DR D. T. ARIYO\t\t\t\t____________________________")
    paragraph(doc, "Head of Department")
    paragraph(doc)
    paragraph(doc, "External Examiner\t\t\t\t____________________________")
    doc.add_page_break()

    heading(doc, "DEDICATION", center=True)
    paragraph(doc, "This project is dedicated to Almighty God for His grace, guidance, strength, and protection throughout the period of this study. It is also dedicated to my family for their encouragement, sacrifice, and continuous support during my academic journey.")
    doc.add_page_break()

    heading(doc, "ACKNOWLEDGEMENT", center=True)
    paragraph(doc, "I sincerely appreciate Almighty God for the knowledge, wisdom, and strength given to me throughout this research work. My profound gratitude goes to my supervisor, Dr. T. A. Adekunle, for his guidance, corrections, patience, and encouragement during the development of this project.")
    paragraph(doc, "I also appreciate the Department of Software Engineering, Faculty of Computing and Information Technology, Osun State University, for providing the academic foundation required to complete this work. I am grateful to my parents, family members, friends, and colleagues for their support, advice, and motivation.")
    doc.add_page_break()

    heading(doc, "TABLE OF CONTENTS", center=True)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    heading(doc, "ABSTRACT", center=True)
    paragraph(doc, f"Cholera remains a major public health problem in Nigeria because outbreaks are associated with poor access to safe water, inadequate sanitation, rainfall, flooding, and delays in disease surveillance reporting. This project developed a predictive model and dashboard for weekly state-level cholera outbreak forecasting in Nigeria. The system was implemented using a modular software pipeline that includes PDF extraction, raw data loading, data cleaning, feature engineering, model training, FastAPI model serving, and React dashboard visualization.")
    paragraph(doc, f"The processed modelling dataset contained {r['rows']} state-period records and {r['columns']} columns across {r['states']} Nigerian states and the Federal Capital Territory from {r['year_min']} to {r['year_max']}. The data recorded {r['cases']:,} suspected cholera cases and {r['deaths']:,} deaths, giving an overall case fatality ratio of {r['cfr'] * 100:.2f}%. Important features such as epidemiological week ranges, reporting gaps, lagged case values, rolling averages, rainfall, flood proxies, WASH indicators, displacement context, health facility counts, and case fatality ratio were engineered from the raw data and API-derived enrichment files.")
    best = r["metrics"][r["metrics"]["model"] == "xgboost"].iloc[0]
    paragraph(doc, f"Five forecasting approaches were compared using expanding time-based cross-validation: naive lag-1 baseline, moving average baseline, Random Forest, XGBoost, and Prophet. XGBoost was selected as the main trainable model because it produced the lowest cross-validation RMSE among the machine learning models, with MAE of {best.mae:.2f}, RMSE of {best.rmse:.2f}, SMAPE of {best.smape:.2f}%, and R2 of {best.r2:.3f}. The system produced four-week forecasts, risk labels, a Nigeria state risk map, case trend chart, forecast chart with uncertainty interval, KPI cards, and a state risk summary. The work demonstrates that software engineering and machine learning can support cholera decision-support systems, although the prototype should not replace official public health judgement or verified surveillance data.")
    doc.add_page_break()

    heading(doc, "LIST OF FIGURES", center=True)
    for item in [
        "Figure 3.1: System architecture of the cholera forecasting platform",
        "Figure 3.2: Use case diagram for the cholera risk platform",
        "Figure 4.1: Nigeria cholera risk map",
        "Figure 4.2: Case trend chart",
        "Figure 4.3: Forecast chart with uncertainty interval",
        "Figure 4.4: State risk summary page",
    ]:
        paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    heading(doc, "LIST OF TABLES", center=True)
    for item in [
        "Table 3.1: Description of major system modules",
        "Table 3.2: Model evaluation metrics",
        "Table 4.1: Dataset summary",
        "Table 4.2: Latest top ten state records by cases",
        "Table 4.3: External enrichment data status",
        "Table 4.4: Model comparison results",
        "Table 4.5: Top forecasted state-weeks",
    ]:
        paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT)


def start_chapters(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_number_format(section, "decimal", 1)
    add_page_number(section, roman=False)


def chapter_one(doc: Document) -> None:
    heading(doc, "CHAPTER ONE", center=True)
    heading(doc, "INTRODUCTION", center=True)
    heading(doc, "1.1 Background of the Study", 2)
    paragraph(doc, "Cholera is an acute diarrhoeal disease caused by the bacterium Vibrio cholerae and is transmitted mainly through food or water contaminated with faecal matter. The World Health Organization describes cholera as a disease that can become fatal within hours if severe dehydration is not treated quickly, while also emphasizing that prevention depends strongly on access to safe water, sanitation, hygiene, surveillance, and rapid response (WHO, 2024). Cholera is therefore not only a medical problem but also a water, sanitation, environmental, and information-management problem.")
    paragraph(doc, "Globally, cholera remains a major public health threat, especially in places affected by weak water infrastructure, population displacement, conflict, flooding, and poor sanitation. WHO reported that cholera cases have continued to rise in recent years, while the ongoing global cholera upsurge has been linked to climate events, conflict, and reduced access to clean water. This makes early warning and surveillance more important because health authorities need timely information to plan treatment centres, distribute oral rehydration supplies, strengthen water sanitation interventions, and monitor high-risk communities.")
    paragraph(doc, "In Nigeria, cholera outbreaks remain recurrent and are documented through situation reports and weekly epidemiological updates published by the Nigeria Centre for Disease Control and Prevention. These reports provide important information on suspected cases, deaths, affected states, and reporting periods. However, public health reports are not always ready for machine learning because they may be stored as PDFs, may use inconsistent epidemiological week ranges, and may summarize several weeks together. Therefore, there is a practical need for a software system that can transform surveillance-style records into a structured modelling dataset.")
    paragraph(doc, "Recent studies support the use of data science and machine learning for cholera prediction. Amshi et al. (2024) developed a cholera outbreak risk prediction model for Nigeria using machine learning and design science principles, showing the usefulness of computational approaches in supporting outbreak detection. Omankwu and Etuk (2024) also proposed a machine-learning approach for early detection and prediction of cholera outbreaks in Nigeria by integrating environmental, socio-economic, and health-related data. These studies show that predictive modelling can support public health planning when it is based on relevant disease and environmental indicators.")
    paragraph(doc, "Environmental studies further show that cholera occurrence is influenced by climate and ecological variables. De Magny et al. (2008) found significant relationships between cholera patterns and environmental signatures such as rainfall and chlorophyll-related indicators, while Daisy et al. (2020) developed a forecasting model in Dhaka using time-series climate data and found rainfall and maximum temperature to be important predictors. In Nigeria, Abdullahi et al. (2023) reported that rainfall and temperature were directly related to cholera occurrence in Kano State. These findings justify the inclusion of rainfall, temperature, humidity, and rainy-season indicators in a cholera forecasting model.")
    paragraph(doc, "This project therefore focuses on the development of a predictive model and dashboard for weekly state-level cholera outbreak forecasting in Nigeria. The system combines raw cholera case reports, deaths, case fatality ratio, epidemiological week information, lagged case features, rolling averages, environmental variables, machine learning models, a FastAPI backend, and a React dashboard. The goal is not to replace public health experts, but to provide an academic decision-support prototype that can summarize patterns and present forecasts in an interpretable form.")
    heading(doc, "1.2 Statement of the Problem", 2)
    paragraph(doc, "Cholera control in Nigeria is affected by the difficulty of detecting risk early enough for preventive action. Surveillance reports are useful, but they are usually reactive because they describe cases that have already occurred. If health authorities only respond after cases have increased, interventions such as water treatment, sanitation campaigns, medical supply distribution, and public warning may be delayed.")
    paragraph(doc, "A second problem is the structure of the available data. Cholera situation reports may be published as PDF documents, and reporting periods are not always consistent. Some reports describe one epidemiological week, while others summarize multi-week ranges such as weeks 1-4 or 6-9. This irregularity makes ordinary time-series modelling more difficult because observations do not always represent equal time intervals. A forecasting system must therefore preserve the original reporting label while extracting start week, end week, and reporting duration.")
    paragraph(doc, "A third problem is that cholera data are usually highly skewed. Many state-period observations may contain low case counts, while a few outbreak spikes may contain very high values. If risk thresholds are calculated from the raw distribution without considering outliers, the resulting low, medium, and high risk labels may not represent the normal data pattern. This problem was also observed in related machine-learning studies where outlier treatment and class imbalance methods were needed to improve outbreak prediction.")
    paragraph(doc, "Finally, several outbreak prediction projects focus either on model development or on visualization, but a useful software engineering project should connect the full workflow. There is a need for a modular system that supports raw data replacement, feature engineering, model comparison, risk classification, API model serving, and dashboard visualization. This study addresses these problems by developing an end-to-end cholera risk prediction platform for Nigeria.")
    heading(doc, "1.3 Aim and Objectives", 2)
    paragraph(doc, "The aim of this study is to develop a predictive model and dashboard for forecasting cholera outbreak risk in Nigeria.")
    for item in [
        "To collect and prepare state-level cholera case and death data for modelling.",
        "To engineer relevant time-based, environmental, lagged, and risk-related features.",
        "To train and compare baseline models, Random Forest, XGBoost, and Prophet.",
        "To classify predicted cases into low, medium, and high risk levels.",
        "To implement a FastAPI backend and React dashboard for model serving and visualization.",
    ]:
        bullet(doc, item)
    heading(doc, "1.4 Research Questions", 2)
    for item in [
        "How can irregular cholera reporting periods be transformed into a structured modelling dataset?",
        "Which algorithm gives the best forecasting performance on the available state-level cholera data?",
        "How can forecasted cases be converted into interpretable risk categories?",
        "How can the results be presented in a dashboard for public health decision support?",
    ]:
        number(doc, item)
    heading(doc, "1.5 Scope of the Study", 2)
    paragraph(doc, "The scope of this study is limited to weekly state-level cholera forecasting in Nigeria. The system uses cases, deaths, case fatality ratio, epidemiological week, reporting gap, lagged cases, rolling averages, rainfall, temperature, humidity, and state boundary data where available. The system predicts short-term case counts and converts the forecasts into low, medium, and high risk categories.")
    paragraph(doc, "The dashboard presents national and state-level outputs through KPI cards, a Nigeria choropleth map, a selected-state trend chart, a forecast chart with uncertainty interval, and a state risk summary page. The study does not perform laboratory confirmation, clinical diagnosis, individual-level patient prediction, or local government area-level forecasting. It also does not replace the authority of NCDC, state ministries of health, WHO, or qualified epidemiologists.")
    heading(doc, "1.6 Significance of the Study", 2)
    paragraph(doc, "This study is significant because it demonstrates how software engineering and machine learning can support disease surveillance. The work is important to public health because it shows how historical outbreak data can be transformed into forecast outputs and visual risk summaries. It is also important academically because it applies data processing, feature engineering, model evaluation, API design, and frontend development to a real health-related problem.")
    paragraph(doc, "The study also contributes to the growing body of literature on computational cholera surveillance in Nigeria. While previous studies have examined machine learning and geospatial methods for cholera prediction, this project emphasizes a practical end-to-end prototype that is suitable for undergraduate software engineering research. It shows how raw NCDC-style reporting data can be connected to model serving and dashboard interpretation.")
    heading(doc, "1.7 Definition of Terms", 2)
    for item in [
        "Epidemiological week: A standardized week used for disease surveillance reporting.",
        "Case fatality ratio: The proportion of deaths among reported cases.",
        "Feature engineering: The process of creating useful model variables from raw data.",
        "Forecasting: The process of predicting future values from historical and recent data.",
        "Risk level: A category that describes whether predicted outbreak risk is low, medium, or high.",
    ]:
        bullet(doc, item)


def chapter_two(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "CHAPTER TWO", center=True)
    heading(doc, "LITERATURE REVIEW", center=True)
    heading(doc, "2.1 Introduction", 2)
    paragraph(doc, "This chapter reviews the concepts, theories, and empirical studies that support the development of a predictive model for cholera outbreaks in Nigeria. The review covers cholera surveillance, environmental determinants, epidemiological reporting, predictive modelling, machine learning, time-series forecasting, geospatial risk mapping, and dashboard-based decision support. The purpose of the review is to show how existing research supports the design choices adopted in this project.")
    heading(doc, "2.2 Cholera and Public Health Surveillance", 2)
    paragraph(doc, "Public health surveillance refers to the continuous collection, analysis, interpretation, and dissemination of health data for action. In cholera control, surveillance helps identify where cases are increasing, where deaths are occurring, and where emergency public health response may be required. According to WHO (2024), countries at risk of cholera require strong epidemiological and laboratory surveillance to quickly detect outbreaks and guide response.")
    paragraph(doc, "Cholera surveillance reports commonly include suspected cases, deaths, affected areas, and reporting periods. In Nigeria, the Nigeria Centre for Disease Control and Prevention publishes disease situation reports that include cholera updates. These reports are valuable because they provide official public health information, but they are not always structured for direct use in machine learning. They may appear as PDFs, contain tables that require extraction, or summarize multiple epidemiological weeks together.")
    paragraph(doc, "For a predictive system, surveillance data must be transformed from report format into an analytical dataset. This requires column normalization, state-name standardization, handling of missing values, conversion of epidemiological week labels, and preservation of reporting intervals. This project therefore treats surveillance data preparation as a major system component rather than a minor preprocessing step.")
    heading(doc, "2.3 Environmental Factors Affecting Cholera", 2)
    paragraph(doc, "Cholera is strongly connected to water, sanitation, and environmental conditions. Heavy rainfall and flooding can contaminate drinking water sources, while drought can reduce access to safe water and increase dependence on unsafe sources. WHO also links cholera risk with limited access to safe water, sanitation, hygiene, displacement, conflict, and climate events such as floods and droughts.")
    paragraph(doc, "Several studies support the inclusion of climate variables in cholera forecasting. De Magny et al. (2008) examined environmental signatures associated with cholera epidemics and reported relationships between cholera patterns and environmental indicators such as rainfall anomalies and aquatic ecosystem conditions. Daisy et al. (2020) developed a forecasting model for cholera incidence in Dhaka using climate time-series data and found rainfall and maximum temperature to be important variables. In another time-series study, Ali et al. (2013) examined cholera incidence in Matlab, Bangladesh, and found that temperature-related variables were useful in explaining cholera incidence.")
    paragraph(doc, "In Nigeria, Abdullahi et al. (2023) assessed the influence of temperature and rainfall on cholera occurrence in Kano State using geospatial methods. Their findings showed that temperature and rainfall were positively related to cholera occurrence, particularly during the warm and wet season. This supports the use of rainfall, temperature, humidity, and rainy-season flags as environmental features in this project.")
    heading(doc, "2.4 Predictive Modelling in Public Health", 2)
    paragraph(doc, "Predictive modelling involves the use of historical data, statistical techniques, and machine learning algorithms to estimate future outcomes. In public health, predictive modelling can support planning by identifying likely increases in disease cases before the burden becomes severe. This can assist in resource allocation, outbreak preparedness, risk communication, and targeted intervention.")
    paragraph(doc, "A disease forecasting model may be designed either as a regression model or as a classification model. A regression model predicts expected case counts, while a classification model predicts categories such as outbreak or non-outbreak. For this project, regression was selected as the main task because predicted case counts are more informative for planning. The predicted cases are then converted into low, medium, and high risk levels for easier interpretation.")
    heading(doc, "2.5 Machine Learning Algorithms", 2)
    paragraph(doc, "Random Forest is an ensemble learning algorithm that combines many decision trees. It is widely used in public health prediction because it can model nonlinear relationships and handle structured datasets. Jutla et al. (2020) demonstrated the potential of Random Forest classifiers and remotely sensed climate variables for environmental cholera risk prediction, showing strong classification performance in coastal India.")
    paragraph(doc, "XGBoost is a gradient boosting algorithm that builds decision trees sequentially and improves predictions by learning from previous errors. It is suitable for tabular data containing lag features, environmental variables, and time-based variables. Amshi et al. (2024) used extreme-gradient boosting in a cholera outbreak risk prediction model for Nigeria and reported strong classification performance after dimensionality reduction, data balancing, and outlier handling.")
    paragraph(doc, "Prophet is a time-series forecasting tool developed to model trend and seasonality. It is useful as a comparison model where disease data have temporal patterns, although it can perform poorly when state-level records are sparse, irregular, or highly volatile. This project includes Prophet for comparison but does not assume it will outperform tree-based models.")
    heading(doc, "2.6 Geospatial Risk Mapping and Dashboards", 2)
    paragraph(doc, "Geospatial analysis is important in cholera surveillance because outbreaks are not evenly distributed across space. Some areas may be more vulnerable due to water sources, sanitation infrastructure, population density, flooding, or movement patterns. Abdullahi et al. (2023) used geospatial methods to examine the relationship between rainfall, temperature, and cholera occurrence in Kano State. Such studies show the importance of mapping outbreak risk in addition to producing numerical predictions.")
    paragraph(doc, "A dashboard helps translate model outputs into visual information. In public health systems, dashboards can display maps, trends, summary indicators, and risk categories. For this project, a React dashboard was selected instead of a Streamlit dashboard because React provides more flexibility for building a structured two-page interface with a Nigeria map, KPI cards, forecast charts, and state risk summary.")
    heading(doc, "2.7 Empirical Review", 2)
    table(doc, "Table 2.1: Summary of related empirical studies", ["Author(s)", "Method", "Major finding", "Relevance to this study"], [
        ["Amshi et al. (2024)", "Machine learning, NMF, SMOTE, DBSCAN, XGBoost", "Machine learning improved cholera outbreak risk prediction in Nigeria.", "Supports model comparison and outlier-aware risk prediction."],
        ["Omankwu and Etuk (2024)", "Machine learning with health, environmental, and socio-economic data", "Proposed data-driven cholera prediction for Nigeria.", "Supports Nigerian cholera ML decision-support direction."],
        ["Jutla et al. (2020)", "Random Forest with essential climate variables", "Climate variables helped identify environmental cholera risk.", "Supports use of environmental variables and Random Forest."],
        ["Daisy et al. (2020)", "SARIMA with rainfall and temperature", "Rainfall and maximum temperature improved cholera forecasting.", "Supports climate-informed time-series forecasting."],
        ["Abdullahi et al. (2023)", "OLS and geographically weighted regression", "Temperature and rainfall influenced cholera occurrence in Kano State.", "Supports geospatial and climate-based Nigerian analysis."],
        ["Magers et al. (2026)", "Climate-driven SIR model", "Climate variables improved cholera prediction for Nigeria and DRC.", "Supports climate-driven modelling for African cholera contexts."],
    ])
    heading(doc, "2.8 Summary of Literature Gap", 2)
    paragraph(doc, "The reviewed literature shows that cholera prediction can benefit from surveillance data, climate variables, outlier handling, geospatial interpretation, and machine learning. However, many studies focus mainly on modelling performance and do not provide a complete software pipeline from raw report data to API model serving and interactive dashboard visualization. Some studies also assume more regular data than situation reports usually provide.")
    paragraph(doc, "This project addresses the identified gap by developing an end-to-end prototype that handles irregular epidemiological week ranges, preserves raw data boundaries, engineers lagged and environmental features, compares baseline and machine learning models, derives robust risk classes, serves outputs through FastAPI, and visualizes them through a React dashboard.")


def chapter_three(doc: Document, r: dict[str, object]) -> None:
    doc.add_page_break()
    heading(doc, "CHAPTER THREE", center=True)
    heading(doc, "METHODOLOGY AND SYSTEM DESIGN", center=True)
    heading(doc, "3.1 Introduction", 2)
    paragraph(doc, "This chapter explains the methodology used in developing the cholera outbreak prediction system. It describes the research design, data sources, data preprocessing, feature engineering, risk threshold design, model development, evaluation strategy, backend API, and dashboard implementation. The methodology follows the direction established in the literature review, where cholera prediction is treated as a data-driven public health decision-support problem involving epidemiological, environmental, temporal, and spatial information.")
    heading(doc, "3.2 Research Design", 2)
    paragraph(doc, "The study adopted an experimental software engineering research design. This design was selected because the project involved building, testing, and evaluating a working software artifact. The artifact consisted of a data pipeline, feature engineering module, forecasting module, model-serving API, and interactive dashboard. This approach is consistent with design-oriented machine learning studies such as Amshi et al. (2024), where the goal is not only to test an algorithm but also to produce a useful decision-support system.")
    paragraph(doc, "The methodological process followed six major stages. First, cholera case and death data were prepared from NCDC-style situation report records. Second, environmental, WASH, flood, displacement, and health-system context variables were collected or prepared from available public APIs and downloadable sources. Third, multiple forecasting models were trained and compared. Fourth, predicted case counts were converted into risk labels. Fifth, the best model and forecast artifacts were exposed through FastAPI. Sixth, the results were visualized in a React dashboard.")
    heading(doc, "3.3 System Architecture", 2)
    paragraph(doc, "The system architecture was designed as a modular pipeline to make the project easier to maintain and extend. The major layers are data collection, data processing, feature engineering, model training, model serving, and visualization. The backend was implemented in Python because of its strong data science ecosystem, while the frontend dashboard was implemented with React to provide a more flexible user interface.")
    paragraph(doc, "A modular design was necessary because manually collected NCDC data may later replace the temporary dataset. Therefore, the raw data file was kept separate from the processed modelling dataset. The system expects raw cholera data in cholera_data.csv and creates the engineered dataset separately. This design prevents feature-engineered columns from being mixed into the raw dataset and supports future data replacement.")
    table(doc, "Table 3.1: Description of major system modules", ["Module", "Description"], [
        ["PDF extraction", "Downloads and extracts NCDC cholera report tables."],
        ["Data pipeline", "Loads raw cholera CSV data, normalizes columns, and merges climate, WASH, flood, displacement, and health facility variables."],
        ["API collectors", "Collects NASA POWER climate, WDI WASH, rainfall flood proxies, UNHCR/DTM displacement context, GRID3 health facilities, and WASHNORM state WASH files where available."],
        ["Feature engineering", "Creates epi-week, lag, rolling, CFR, climate, WASH, flood proxy, displacement, health facility, and risk label features."],
        ["Model training", "Compares baseline, Random Forest, XGBoost, and Prophet models."],
        ["FastAPI backend", "Serves model, history, forecast, boundary, and prediction endpoints."],
        ["React dashboard", "Displays KPIs, map, trend, forecast, and state risk summary."],
    ])
    figure(doc, "figure_3_1_architecture", "System architecture screenshot or diagram will be inserted here.", "Figure 3.1: System architecture of the cholera forecasting platform.")
    heading(doc, "3.4 Data Collection", 2)
    paragraph(doc, f"The active processed dataset used for this project contained {r['rows']} state-period observations across {r['states']} Nigerian states and the Federal Capital Territory from {r['year_min']} to {r['year_max']}. The raw input file is designed to be replaced with manually collected NCDC data using the same filename, cholera_data.csv. This makes the system suitable for continuous improvement when more verified surveillance records are available.")
    paragraph(doc, "The expected raw data fields include state, year, epidemiological week, cases, deaths, and optional case fatality ratio. The data structure accepts both single epidemiological weeks and multi-week reporting ranges. This is important because real situation reports may describe periods such as week 1, week 1-4, or week 6-9. The system therefore treats epidemiological week parsing as a core preprocessing step.")
    paragraph(doc, "A major methodological challenge was that cholera outbreak situation reports were not uploaded consistently every week. Some reporting periods were missing, some reports were released after delays, and some reports summarized several epidemiological weeks at once instead of providing one record for every week. As a result, the dataset is not a perfectly regular weekly time series. Building a model from this type of surveillance data is complex because the model must learn from uneven reporting intervals, multi-week aggregates, and gaps between reports. For this reason, the pipeline preserved the original epi-week label, extracted epi-week start and end values, calculated the number of weeks covered by each report, and added a reporting-gap feature.")
    paragraph(doc, "Environmental data were considered because the literature shows strong links between cholera and climate variables. The system collected rainfall, temperature, and humidity values from NASA POWER for Nigerian state centroids. From the weekly rainfall data, additional flood proxy variables were engineered, including rainfall anomaly, rainfall z-score, heavy-rain flag, extreme-rain flag, and four-week rolling rainfall. ReliefWeb flood report counts were also implemented as an optional API-derived feature, but the local request returned a 403 access response during the latest run. Therefore, the current flood feature file was generated from rainfall proxies while the ReliefWeb report-count column was retained as zero.")
    paragraph(doc, "The system also added contextual predictors beyond climate. World Bank WDI was used to collect national annual WASH indicators, including basic water access, safely managed water access, basic sanitation, safely managed sanitation, and open defecation. A WASHNORM state-profile downloader and extractor was implemented for state-level WASH data, but the UNICEF page blocked direct automated requests during the latest run. Therefore, the state WASH file currently exists as an empty schema, while WDI WASH remains the active fallback. Displacement context was implemented through a DTM path when a subscription key is available and a public UNHCR fallback when no DTM key is supplied. Health-system context was collected from the public GRID3/NHFR-derived ArcGIS service, which produced state-level total facility counts, hospital counts, and primary health care counts.")
    heading(doc, "3.5 Feature Engineering", 2)
    paragraph(doc, "Feature engineering was performed to convert raw surveillance records into machine-learning-ready variables. The system extracted epidemiological week start and end values from both single-week and multi-week reporting periods. It also calculated the number of weeks covered by each report and the reporting gap between consecutive reports for the same state.")
    paragraph(doc, "Time-based variables included date, month, quarter, epidemiological week, and rainy-season indicator. These variables were included because cholera occurrence may follow seasonal patterns related to rainfall, flooding, and water contamination. Environmental variables included rainfall, temperature, humidity, rainfall anomaly, rainfall z-score, heavy-rain flag, extreme-rain flag, and rolling four-week rainfall. Context variables included WASH percentages, displacement population proxies, displacement data age, health facility count, hospital count, and PHC count. State codes were added to provide a numeric representation of state identity for machine learning models.")
    paragraph(doc, "Lagged variables were also engineered because disease counts are often dependent on recent history. The model used lag-1, lag-2, lag-4, and lag-8 case values, as well as lagged deaths and lagged CFR values. Rolling averages were added to smooth short-term variation and represent recent outbreak momentum. This is consistent with time-series disease forecasting studies that use autoregressive and lagged structures to capture temporal dependence.")
    heading(doc, "3.6 Risk Classification", 2)
    paragraph(doc, "Risk classification was required because public health users may find low, medium, and high categories easier to interpret than raw predicted case counts alone. The system first predicts case counts as a regression problem, then converts the predicted counts into risk levels. This approach preserves the numerical forecast while also providing a simple category for dashboard interpretation.")
    paragraph(doc, "The risk thresholds were derived after handling outliers. Cholera data often contain sudden spikes that can distort raw quantiles. To reduce this effect, the system calculates the interquartile range and removes values above the upper fence before computing the median and seventy-fifth percentile. The trimmed median serves as the low-to-medium threshold, while the trimmed seventy-fifth percentile serves as the medium-to-high threshold. This outlier-aware approach is consistent with the need for outlier handling reported in Nigerian machine-learning studies such as Amshi et al. (2024).")
    heading(doc, "3.7 Model Training and Evaluation", 2)
    paragraph(doc, "The models compared were naive lag-1 baseline, moving average baseline, Random Forest, XGBoost, and Prophet. The baseline models were included to provide simple reference points. Random Forest and XGBoost were selected because the literature shows that tree-based methods are useful for structured disease and environmental datasets. Prophet was included as a time-series comparison method because cholera data may contain trend and seasonal behaviour.")
    paragraph(doc, "Since the data are time ordered, random train-test splitting was not used. Random splitting can leak future information into the training set and make performance appear better than it really is. Instead, the data were sorted chronologically from 2021 to 2025. Model comparison used expanding time-based cross-validation, where the training window expands forward in time and each test fold represents later reporting periods. The latest holdout split also ensured that training periods ended before the test periods began.")
    paragraph(doc, "The evaluation metrics were selected to measure different aspects of forecasting performance. Mean Absolute Error measured the average size of prediction errors. Root Mean Squared Error penalized large errors more strongly. SMAPE provided a percentage-based error measure. R2 measured the degree to which the model explained variation in the target variable.")
    table(doc, "Table 3.2: Model evaluation metrics", ["Metric", "Meaning"], [
        ["MAE", "Average absolute difference between actual and predicted cases."],
        ["RMSE", "Square-root of average squared prediction error."],
        ["SMAPE", "Percentage error metric that is more stable for varying case values."],
        ["R2", "Proportion of variation explained by the model."],
    ])
    heading(doc, "3.8 Backend API Implementation", 2)
    paragraph(doc, "The backend was implemented with FastAPI. FastAPI was selected because it supports clear API route definitions, automatic validation, and easy integration with Python model artifacts. The API includes routes for health checking, model information, metrics, summary data, historical records, forecasts, state boundaries, single prediction, batch prediction, and artifact reload. These routes separate model-serving logic from dashboard presentation logic.")
    paragraph(doc, "The forecast endpoint returns one-to-four-week recursive forecasts for each state. The summary and history endpoints provide dashboard data such as cases, deaths, CFR, risk levels, and climate variables. The boundaries endpoint serves the Nigeria state boundary GeoJSON used by the map. This API design allows the frontend to remain independent of the modelling code.")
    heading(doc, "3.9 Dashboard Implementation", 2)
    paragraph(doc, "The frontend was implemented using React, Vite, and Leaflet. React was selected because it supports component-based dashboard development. Leaflet was used to display the Nigeria state choropleth map from GeoJSON boundaries. The dashboard has two pages. The Overview page shows KPI cards, a Nigeria choropleth map, selected-state details, a case trend chart, and a forecast chart with uncertainty interval. The State Risk Summary page shows the top ten states for the selected reporting period.")
    paragraph(doc, "The dashboard includes an epidemiological-week filter at the top-right corner. When a reporting period is selected, the KPI cards, map, and state risk summary update accordingly. Clicking a state on the map selects that state and updates the forecast and trend charts. This design supports interactive exploration while keeping model performance tables off the dashboard, since the dashboard is intended for interpretation rather than technical model evaluation.")
    figure(doc, "figure_3_2_use_case", "Use case diagram screenshot will be inserted here.", "Figure 3.2: Use case diagram for the cholera risk platform.")


def chapter_four(doc: Document, r: dict[str, object]) -> None:
    doc.add_page_break()
    data = r["data"]
    metrics = r["metrics"]
    validation = r["validation"]
    forecast = r["forecast"]
    latest = r["latest"]
    heading(doc, "CHAPTER FOUR", center=True)
    heading(doc, "RESULTS AND DISCUSSION", center=True)
    heading(doc, "4.1 Introduction", 2)
    paragraph(doc, "This chapter presents the system implementation results, dataset summary, model comparison results, validation output, forecast output, dashboard output, and discussion of findings.")
    heading(doc, "4.2 Dataset Results", 2)
    paragraph(doc, f"The processed dataset contained {r['rows']} observations, {r['columns']} columns, {r['states']} states, and {r['periods']} unique reporting periods. The data covered {r['year_min']} to {r['year_max']} and recorded {r['cases']:,} suspected cases and {r['deaths']:,} deaths. The overall CFR was {r['cfr'] * 100:.2f}%.")
    table(doc, "Table 4.1: Dataset summary", ["Statistic", "Value"], [
        ["Processed observations", f"{r['rows']:,}"],
        ["Processed columns", f"{r['columns']:,}"],
        ["States covered", r["states"]],
        ["Unique reporting periods", r["periods"]],
        ["Year range", f"{r['year_min']} - {r['year_max']}"],
        ["Total suspected cases", f"{r['cases']:,}"],
        ["Total deaths", f"{r['deaths']:,}"],
        ["Overall CFR", f"{r['cfr'] * 100:.2f}%"],
    ])
    risk_counts = data["risk_level"].value_counts().to_dict()
    paragraph(doc, f"The full dataset contained {risk_counts.get('Low', 0)} low-risk records, {risk_counts.get('Medium', 0)} medium-risk records, and {risk_counts.get('High', 0)} high-risk records.")
    top_latest = latest.sort_values("suspected_cases", ascending=False).head(10)
    table(doc, "Table 4.2: Latest top ten state records by cases", ["State", "Year", "Epi week", "Cases", "Deaths", "CFR", "Risk"], [
        [row.state, int(row.year), row.epi_week_label, int(row.suspected_cases), int(row.deaths), f"{row.cfr * 100:.2f}%", row.risk_level]
        for row in top_latest.itertuples()
    ])
    figure(doc, "figure_4_1_map", "Nigeria risk map screenshot will be inserted here.", "Figure 4.1: Nigeria cholera risk map.")
    heading(doc, "4.3 External Enrichment Data Status", 2)
    paragraph(doc, "The current implementation includes several API-derived or download-derived enrichment files. The status of each source as at the latest project run is shown in Table 4.3.")
    table(doc, "Table 4.3: External enrichment data status", ["Source", "Output file", "Rows", "Current status"], [
        [
            item["source"],
            item["file"],
            item["rows"],
            "Available" if item["rows"] > 0 else "Schema/fallback only",
        ]
        for item in r["enrichment_status"]
    ])
    paragraph(doc, "NASA POWER climate, rainfall flood proxies, World Bank WDI WASH, UNHCR displacement fallback, and GRID3 health facility features were available during the latest build. ReliefWeb flood counts were implemented but the request returned a 403 access response, so the rainfall-based flood proxy columns were used and the report-count feature remained zero. UNICEF WASHNORM direct scraping also returned a 403 response, so state-level WASH values were not available through that automated route in the latest run.")
    heading(doc, "4.4 Model Comparison Results", 2)
    paragraph(doc, "The model comparison was based on three expanding time-based cross-validation folds. This ensured that future periods were not used to train predictions for earlier periods.")
    table(doc, "Table 4.4: Model comparison results", ["Model", "Status", "Folds", "MAE", "RMSE", "SMAPE", "R2"], [
        [row.model, row.status, int(row.folds), f"{row.mae:.2f}", f"{row.rmse:.2f}", f"{row.smape:.2f}%", f"{row.r2:.3f}"]
        for row in metrics.itertuples()
    ])
    best = metrics[metrics["model"] == "xgboost"].iloc[0]
    rf = metrics[metrics["model"] == "random_forest"].iloc[0]
    paragraph(doc, f"XGBoost was selected as the final trainable model because it produced the lowest RMSE among the machine learning models, with MAE of {best.mae:.2f}, RMSE of {best.rmse:.2f}, SMAPE of {best.smape:.2f}%, and R2 of {best.r2:.3f}. Random Forest recorded a slightly lower MAE of {rf.mae:.2f}, but its RMSE was higher at {rf.rmse:.2f}. Prophet performed poorly on the current irregular state-level dataset, with RMSE of {metrics[metrics['model'] == 'prophet'].iloc[0].rmse:.2f}.")
    heading(doc, "4.5 Validation Results", 2)
    max_error = validation.sort_values("absolute_error", ascending=False).iloc[0]
    paragraph(doc, f"The latest holdout validation set contained {len(validation)} observations. The mean actual case count was {validation['actual_cases'].mean():.2f}, the mean predicted case count was {validation['predicted_cases'].mean():.2f}, and the mean absolute validation error was {validation['absolute_error'].mean():.2f}. The highest validation error occurred in {max_error.state} in {int(max_error.year)} week {int(max_error.epi_week)}, where the actual cases were {max_error.actual_cases:.0f} and the predicted cases were {max_error.predicted_cases:.2f}.")
    figure(doc, "figure_4_2_trend", "Case trend screenshot will be inserted here.", "Figure 4.2: Case trend chart.")
    heading(doc, "4.6 Forecast Results", 2)
    forecast_counts = forecast["risk_level"].value_counts().to_dict()
    paragraph(doc, f"The final model generated {len(forecast)} forecast records representing four forecast weeks for each state. The forecast output contained {forecast_counts.get('Low', 0)} low-risk forecasts, {forecast_counts.get('Medium', 0)} medium-risk forecasts, and {forecast_counts.get('High', 0)} high-risk forecasts.")
    top_forecast = forecast.sort_values("predicted_cases", ascending=False).head(10)
    table(doc, "Table 4.5: Top forecasted state-weeks", ["State", "Forecast week", "Year", "Epi week", "Predicted cases", "Lower", "Upper", "Risk"], [
        [row.state, int(row.forecast_week), int(row.year), int(row.epi_week), f"{row.predicted_cases:.2f}", f"{row.predicted_lower:.2f}", f"{row.predicted_upper:.2f}", row.risk_level]
        for row in top_forecast.itertuples()
    ])
    figure(doc, "figure_4_3_forecast", "Forecast chart with uncertainty interval screenshot will be inserted here.", "Figure 4.3: Forecast chart with uncertainty interval.")
    heading(doc, "4.7 Dashboard Results", 2)
    paragraph(doc, "The dashboard presented model outputs in a user-friendly format. The Overview page displayed cases, deaths, CFR, a Nigeria map, selected-state details, a case trend chart, and a forecast chart. The second page displayed the State Risk Summary table showing the top ten states by risk and cases for the selected reporting period.")
    figure(doc, "figure_4_4_summary", "State risk summary screenshot will be inserted here.", "Figure 4.4: State risk summary page.")
    heading(doc, "4.8 Discussion", 2)
    paragraph(doc, "The results show that tree-based models were more suitable for the available structured dataset than Prophet. XGBoost produced the best overall result among the trainable models after randomized hyperparameter tuning. After adding the currently available enrichment variables, the R2 score became slightly positive at 0.001. This means that the model performed only marginally better than predicting the average case count, but it also shows that the enriched pipeline improved the earlier negative R2 result.")
    paragraph(doc, "There are several reasons for the low R2 score. First, cholera outbreak data are highly irregular and contain sudden spikes. For example, a state may report only a few cases in one period and then report a very large outbreak in the next available report. These abrupt increases are difficult for a model to predict unless the triggering factors are available in the data before the spike occurs.")
    paragraph(doc, "Second, the situation reports used for the dataset were not uploaded consistently every week. Some reports covered one epidemiological week, while others covered three or four weeks together. Some weeks were also missing entirely. This means the model was not learning from a continuous weekly time series, but from uneven surveillance records. Even though reporting-gap and period-length features were added, irregular reporting still reduces the ability of the model to explain case variation.")
    paragraph(doc, "Third, the dataset was state-level rather than community-level or local government-level. Cholera outbreaks often begin in specific communities, water sources, camps, or local government areas before becoming visible at state level. Aggregating the data to state level can hide local outbreak signals. This can reduce R2 because the model receives broad state summaries instead of detailed early-warning indicators.")
    paragraph(doc, "Fourth, some important predictors are still only partially represented. The current dataset now includes WDI WASH indicators, NASA climate variables, rainfall-derived flood proxies, a UNHCR displacement fallback, and GRID3 health facility counts. However, the strongest versions of these predictors were not fully available through unauthenticated automated collection. ReliefWeb flood counts were blocked by a 403 response, UNICEF WASHNORM state-profile scraping was blocked by a 403 response, and true DTM state-level displacement data require a subscription key. Therefore, the model still lacks precise local WASH conditions, confirmed flood-event exposure, local displacement movements, vaccination coverage, intervention timing, and health facility reporting-delay variables.")
    paragraph(doc, "Finally, the current dataset is still limited in size for a state-level machine learning task. Although it contains several hundred processed observations, the number of observations per state is relatively small after feature engineering and chronological splitting. This limits the ability of complex models to generalize. Therefore, the low R2 score should not be interpreted as a total failure of the system. Instead, it shows that the current data are noisy, irregular, and incomplete, and that the prototype is more reliable as a decision-support and visualization system than as a fully operational outbreak prediction engine.")


def chapter_five(doc: Document, r: dict[str, object]) -> None:
    doc.add_page_break()
    data = r["data"]
    metrics = r["metrics"]
    best = metrics[metrics["model"] == "xgboost"].iloc[0]
    heading(doc, "CHAPTER FIVE", center=True)
    heading(doc, "SUMMARY, CONCLUSION AND RECOMMENDATIONS", center=True)
    heading(doc, "5.1 Summary", 2)
    paragraph(doc, "This chapter presents the summary, conclusion, contributions, limitations, recommendations, and suggestions for further study. The aim of the project was to develop a predictive model and dashboard for cholera outbreak forecasting in Nigeria using state-level surveillance data and relevant contextual variables. The work was motivated by the need for a software-based decision-support tool that can organize irregular cholera situation report data, produce forecasts, classify risk, and present results in a format that is easier to interpret.")
    paragraph(doc, f"The project produced a complete software prototype consisting of a data pipeline, feature engineering layer, model training module, FastAPI backend, and React dashboard. The processed modelling dataset contained {r['rows']} observations and {r['columns']} columns across {r['states']} states from {r['year_min']} to {r['year_max']}. The dataset recorded {r['cases']:,} suspected cholera cases and {r['deaths']:,} deaths, producing an overall case fatality ratio of {r['cfr'] * 100:.2f}%. The raw data structure was designed so that the temporary dataset can later be replaced with manually collected NCDC data using the same cholera_data.csv filename.")
    paragraph(doc, "The system handled one of the major difficulties in the dataset: irregular epidemiological week reporting. Some records covered a single week, some covered three or four weeks, and some weeks were missing entirely. To address this, the feature engineering layer preserved the original epi-week label, extracted start and end weeks, calculated reporting period length, and created a reporting-gap feature. Lagged case values, lagged deaths, lagged CFR, rolling averages, seasonal variables, and risk labels were also generated.")
    paragraph(doc, "In addition to cholera cases and deaths, the implementation collected and merged available external predictors. NASA POWER supplied rainfall, temperature, and humidity. Rainfall was further transformed into flood proxy variables, including rainfall anomaly, rainfall z-score, heavy-rain flag, extreme-rain flag, and four-week rolling rainfall. World Bank WDI supplied national WASH indicators. UNHCR data provided a displacement fallback, while GRID3/NHFR-derived data supplied health facility, hospital, and PHC counts. ReliefWeb flood counts and UNICEF WASHNORM state profiles were implemented as data-source paths, although both returned access restrictions during the latest automated run.")
    paragraph(doc, f"Five modelling approaches were compared using expanding time-based cross-validation: naive lag-1 baseline, moving average baseline, Random Forest, XGBoost, and Prophet. XGBoost was selected as the final model because it produced the best RMSE among the machine learning models, with MAE of {best.mae:.2f}, RMSE of {best.rmse:.2f}, SMAPE of {best.smape:.2f}%, and R2 of {best.r2:.3f}. The model result shows that the prototype can generate structured forecasts, but the very low R2 also shows that state-level cholera prediction remains difficult when outbreak spikes are sudden and key local predictors are incomplete.")
    heading(doc, "5.2 Conclusion", 2)
    paragraph(doc, "The project achieved its main objective by developing a functional cholera risk prediction and visualization platform. The system successfully converts raw state-level cholera records into a modelling dataset, enriches the records with available contextual variables, trains and compares multiple algorithms, generates risk classifications, exposes results through API routes, and presents outputs through an interactive dashboard.")
    paragraph(doc, "The findings show that tree-based machine learning models are more suitable for the current structured dataset than Prophet. XGBoost performed best by RMSE, while Random Forest also produced competitive results. However, the low R2 value confirms that the model explains only a very small part of the variation in future case counts. This does not make the project unsuccessful; rather, it highlights the reality of cholera outbreak modelling with irregular public health reports. Sudden outbreak spikes are difficult to predict without precise local information on WASH conditions, flooding, displacement, interventions, and reporting delays.")
    paragraph(doc, "The final system should therefore be understood as a decision-support prototype rather than a replacement for official epidemiological judgement. Its value lies in organizing data, making risk patterns visible, creating reproducible forecasts, and showing how additional data sources can be integrated into future cholera surveillance systems. With more complete and validated NCDC records, stronger state-level WASH data, better displacement data, and confirmed flood-event indicators, the model can be improved further.")
    heading(doc, "5.3 Contributions of the Study", 2)
    paragraph(doc, "The study made both software engineering and applied machine learning contributions. The main contributions are as follows:")
    for item in [
        "A modular data pipeline that separates raw NCDC-style cholera data from the processed modelling dataset.",
        "A feature engineering approach for inconsistent epi-week reports, including multi-week period parsing, reporting-gap extraction, lagged features, rolling averages, CFR features, and risk labels.",
        "API and download collectors for climate, rainfall flood proxies, WASH indicators, displacement context, and health facility context.",
        "A model comparison workflow using naive baseline, moving average baseline, Random Forest, XGBoost, and Prophet.",
        "A time-aware validation strategy using expanding cross-validation instead of random train-test splitting, reducing the risk of future data leakage.",
        "FastAPI routes for model serving, forecast retrieval, historical data access, dashboard summaries, state boundaries, and prediction requests.",
        "A React dashboard with KPI cards, year filtering, epi-week filtering, Nigeria state map, state labels, trend chart, forecast interval chart, and state risk summary page.",
        "A reproducible project structure that allows the synthetic or temporary dataset to be replaced later by manually collected NCDC data without changing the pipeline design.",
    ]:
        bullet(doc, item)
    heading(doc, "5.4 Limitations of the Study", 2)
    paragraph(doc, "Although the project produced a working prototype, some limitations affected the modelling results and should be considered when interpreting the system outputs.")
    for item in [
        "The current dataset is based on state-level records, which can hide outbreak signals that begin at local government, ward, community, camp, or water-source level.",
        "The situation reports are irregular. Some reporting periods are missing and some records summarize several weeks, making the data less suitable for standard weekly time-series modelling.",
        "Some important predictors are represented only by proxies. WDI WASH is national annual data, UNHCR displacement fallback is not as strong as DTM state-level displacement, and rainfall-derived flood proxies are not the same as confirmed flood-event data.",
        "Automated access to some useful sources was restricted during implementation. ReliefWeb flood counts returned a 403 response, and UNICEF WASHNORM state-profile scraping also returned a 403 response.",
        "The model predicts state-level case counts, but real public health response may require LGA-level or community-level decisions.",
        "The low R2 indicates that the current model still has limited explanatory power for sudden outbreak spikes, even though it can support exploratory forecasting and visualization.",
    ]:
        bullet(doc, item)
    heading(doc, "5.5 Recommendations", 2)
    paragraph(doc, "Based on the implementation and results, the following recommendations are made:")
    for item in [
        "Manually validated NCDC situation report data should be used to replace the temporary dataset before the system is used for serious analysis.",
        "State-level and LGA-level WASH data should be added, especially indicators on safe water access, sanitation access, open defecation, and handwashing facilities.",
        "A valid DTM subscription key should be used to collect Admin 1 displacement data, and displacement should be linked to the reporting date so that future information is not leaked into earlier predictions.",
        "A reliable flood-event source should be integrated. If ReliefWeb access remains blocked, another open flood-event feed or manually validated flood incident file should be used.",
        "The model should be retrained regularly as new surveillance records are added, and each retraining run should save metrics, validation predictions, and forecast artifacts.",
        "The dashboard should be reviewed by public health users to confirm whether the map, filters, KPI cards, forecasts, and state risk summary are understandable and useful for decision support.",
        "Forecast outputs should always be interpreted together with epidemiological judgement, field reports, and official NCDC guidance.",
    ]:
        bullet(doc, item)
    heading(doc, "5.6 Suggestions for Further Study", 2)
    paragraph(doc, "Further studies can extend this work in several directions. First, future researchers can develop an LGA-level cholera forecasting model if more granular data become available. This would make the system more useful for targeted public health response because cholera outbreaks are often localized before they become visible at state level.")
    paragraph(doc, "Second, future work can explore probabilistic forecasting methods that estimate uncertainty more formally. Instead of producing only a point forecast, the system can estimate the probability that a state will exceed a low, medium, or high-risk threshold. This would support better planning under uncertainty.")
    paragraph(doc, "Third, spatial hotspot modelling can be added using neighbouring-state features, distance to previous outbreaks, riverine or coastal exposure, and population movement. This would help the model consider the geographic spread of cholera risk rather than treating each state as completely independent.")
    paragraph(doc, "Fourth, the platform can be expanded to use forecasted rainfall and flood warnings instead of only historical rainfall. This would improve the practical value of the system for early warning because public health officials need information before a reporting period ends.")
    paragraph(doc, "Finally, future studies can improve automated PDF extraction from NCDC situation reports. Better extraction would reduce manual data entry and make it easier to update the system whenever new situation reports are released.")


def references(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "REFERENCES", center=True)
    for ref in [
        "Abdullahi, A. H., Hassan, M. B., Kanoma, M. S., and Yahuza, Y. (2023). The geo-spatial assessment of the influence of temperature and rainfall on the occurrence of cholera in Kano State, Nigeria. FUDMA Journal of Sciences. https://doi.org/10.33003/fjs-2023-0705-1989",
        "Ali, M., Kim, D. R., Yunus, M., and Emch, M. (2013). Time series analysis of cholera in Matlab, Bangladesh, during 1988-2001. Journal of Health, Population and Nutrition, 31(1), 11-19. https://doi.org/10.3329/jhpn.v31i1.14744",
        "Ali, M., Nelson, A. R., Lopez, A. L., and Sack, D. A. (2015). Updated global burden of cholera in endemic countries. PLoS Neglected Tropical Diseases, 9(6), e0003832. https://doi.org/10.1371/journal.pntd.0003832",
        "Amshi, H. A., Prasad, R., Sharma, B. K., Yusuf, S. I., and Sani, Z. (2024). How can machine learning predict cholera: Insights from experiments and design science for action research. Journal of Water and Health, 22(1), 21-35. https://doi.org/10.2166/wh.2023.026",
        "Chen, T., and Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining.",
        "Daisy, S. S., Islam, A. K. M. S., Akanda, A. S., Faruque, A. S. G., Amin, N., and Jensen, P. K. M. (2020). Developing a forecasting model for cholera incidence in Dhaka megacity through time series climate data. Journal of Water and Health, 18(2), 207-223. https://doi.org/10.2166/wh.2020.133",
        "De Magny, G. C., Murtugudde, R., Sapiano, M. R. P., Nizam, A., Brown, C. W., Busalacchi, A. J., Yunus, M., Nair, G. B., Gil, A. I., Lanata, C. F., Calkins, J., Manna, B., Rajendran, K., Bhattacharya, M. K., Huq, A., Sack, R. B., and Colwell, R. R. (2008). Environmental signatures associated with cholera epidemics. Proceedings of the National Academy of Sciences, 105(46), 17676-17681.",
        "Federal Ministry of Health Nigeria. Nigeria Health Facility Registry. https://hfr.fmohconnect.gov.ng/",
        "GRID3. Nigeria health facility data and geospatial resources. https://grid3.org/",
        "International Organization for Migration. Displacement Tracking Matrix API. https://dtm.iom.int/data-and-analysis/dtm-api",
        "Jutla, A., Akanda, A. S., and Islam, S. (2020). Cholera risk: A machine learning approach applied to essential climate variables. International Journal of Environmental Research and Public Health, 17(24), 9378.",
        "Magers, B., and colleagues. (2026). Evaluation of a novel climate-driven SIR model for cholera prediction. GeoHealth. https://doi.org/10.1029/2025GH001437",
        "NASA POWER Project. POWER Data Access Viewer and API documentation. https://power.larc.nasa.gov/",
        "Nigeria Centre for Disease Control and Prevention. Disease situation reports: Cholera outbreak in Nigeria. https://www.ncdc.gov.ng/diseases/sitreps",
        "Omankwu, O. C. B., and Etuk, E. (2024). Leveraging machine learning for early detection and prediction of cholera outbreaks in Nigeria: A data-driven approach. Transactions of the Nigerian Association of Mathematical Physics. https://doi.org/10.60787/tnamp.v20.383",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., and Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "ReliefWeb. ReliefWeb API documentation. https://apidoc.reliefweb.int/",
        "Taylor, S. J., and Letham, B. (2018). Forecasting at scale. The American Statistician, 72(1), 37-45.",
        "UNHCR Operational Data Portal. Nigeria population data. https://data.unhcr.org/en/country/nga",
        "UNICEF Nigeria. WASHNORM 2021 report. https://www.unicef.org/nigeria/reports/water-sanitation-and-hygiene-national-outcome-routine-mapping-report-2021",
        "World Health Organization. (2024). Cholera fact sheet. https://www.who.int/news-room/fact-sheets/detail/cholera",
        "World Health Organization. (2026). Cholera upsurge, 2021-present. https://www.who.int/emergencies/situations/cholera-upsurge",
        "World Bank. World Bank Indicators API documentation. https://datahelpdesk.worldbank.org/knowledgebase/articles/889392-about-the-indicators-api-documentation",
        "World Bank. World Development Indicators. https://databank.worldbank.org/source/world-development-indicators",
        "XGBoost Developers. XGBoost Python API documentation. https://xgboost.readthedocs.io/",
        "Meta Open Source. Prophet forecasting documentation. https://facebook.github.io/prophet/docs/",
    ]:
        paragraph(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT)


def main() -> None:
    r = collect()
    doc = Document()
    configure_styles(doc)
    set_update_fields_on_open(doc)
    add_front_matter(doc, r)
    start_chapters(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc, r)
    chapter_four(doc, r)
    chapter_five(doc, r)
    references(doc)
    doc.save(DOC_PATH)
    print(f"Wrote {DOC_PATH}")


if __name__ == "__main__":
    main()
